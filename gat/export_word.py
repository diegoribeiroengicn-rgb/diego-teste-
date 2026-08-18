"""
Exportação de relatórios e OPRs em Microsoft Word (.docx) totalmente
editável — títulos, subtítulos, textos, observações, indicadores, tabelas,
cabeçalhos, rodapés e conclusões são inseridos como conteúdo real do
documento (parágrafos e tabelas nativas do Word), nunca como imagem ou
captura de tela.

Gráficos são a única exceção: são inseridos como imagens individuais em
alta resolução (renderizadas via Kaleido a partir da mesma figura Plotly
usada na tela, reaproveitando o Chromium já instalado no ambiente), cada
uma com título acima e legenda abaixo, podendo ser removida, movida ou
substituída no Word sem afetar o restante do documento — nunca o
documento inteiro é gerado como uma imagem única.
"""

from __future__ import annotations

import io
import logging
import os
import re
import shutil

import pandas as pd
import plotly.graph_objects as go
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from gat.config import LOGO_PATH
from gat.horario import agora_br

_logger = logging.getLogger(__name__)

_NAVY = RGBColor(0x1B, 0x3A, 0x8A)
_TEXTO_FRACO = RGBColor(0x64, 0x74, 0x8B)
_ESTILO_TABELA = "Light Grid Accent 1"


def _detectar_chrome_path() -> str | None:
    """Localiza um executável de Chrome/Chromium utilizável pelo Kaleido,
    sem assumir um caminho fixo — em ambientes onde o Chromium do
    Playwright não está provisionado (ex.: Streamlit Cloud), um caminho
    fixo inválido faria o Kaleido falhar com `ChromeNotFoundError`.
    Retorna `None` quando nada é encontrado, deixando o Kaleido tentar
    sua própria detecção padrão."""
    candidatos = [os.environ.get("BROWSER_PATH"), os.environ.get("CHROME_PATH")]
    for nome in ("google-chrome", "google-chrome-stable", "chromium", "chromium-browser"):
        candidatos.append(shutil.which(nome))
    candidatos.append("/opt/pw-browsers/chromium")  # Chromium do Playwright, quando provisionado
    for caminho in candidatos:
        if caminho and os.path.isfile(caminho) and os.access(caminho, os.X_OK):
            return caminho
    return None


# Resolvido uma única vez na importação do módulo — reavaliar a cada
# gráfico seria desnecessário, já que o ambiente não muda em tempo de execução.
_KALEIDO_CHROME_PATH = _detectar_chrome_path()


def figura_para_imagem(fig: go.Figure, largura_px: int = 1400, altura_px: int = 800, escala: float = 2.0) -> bytes:
    """Renderiza uma figura Plotly como PNG em alta resolução, para
    inserção individual no Word (nunca uma captura da tela inteira)."""
    import kaleido

    opts = {"width": largura_px, "height": altura_px, "scale": escala}
    kopts = {"path": _KALEIDO_CHROME_PATH} if _KALEIDO_CHROME_PATH else {}
    return kaleido.calc_fig_sync(fig, opts=opts, kopts=kopts)


def novo_documento() -> Document:
    """Cria um documento Word em branco com a formatação-base institucional
    (fonte, margens) — todo o conteúdo adicionado a partir daqui continua
    integralmente editável pelo usuário final."""
    doc = Document()
    estilo_normal = doc.styles["Normal"]
    estilo_normal.font.name = "Calibri"
    estilo_normal.font.size = Pt(10)
    estilo_normal.paragraph_format.space_after = Pt(3)
    for nivel_titulo in (1, 2, 3):
        estilo_titulo = doc.styles[f"Heading {nivel_titulo}"]
        estilo_titulo.paragraph_format.space_before = Pt(6)
        estilo_titulo.paragraph_format.space_after = Pt(3)
    for secao_doc in doc.sections:
        secao_doc.left_margin = Cm(1.8)
        secao_doc.right_margin = Cm(1.8)
        secao_doc.top_margin = Cm(1.3)
        secao_doc.bottom_margin = Cm(1.3)
    return doc


def _paragrafo(doc: Document, texto: str, cor: RGBColor | None = None, tamanho: float | None = None, negrito: bool = False):
    p = doc.add_paragraph()
    run = p.add_run(texto)
    if cor is not None:
        run.font.color.rgb = cor
    if tamanho is not None:
        run.font.size = Pt(tamanho)
    run.font.bold = negrito
    return p


def cabecalho_institucional(
    doc: Document,
    titulo: str,
    subtitulo: str,
    periodo_label: str,
    filtros_aplicados: dict[str, str] | None,
    usuario_responsavel: str,
    compacto: bool = False,
) -> None:
    """Cabeçalho institucional do relatório: logo, título, subtítulo,
    período selecionado, filtros aplicados, data/hora de geração e usuário
    responsável — cada um como parágrafo próprio, editável individualmente.
    `compacto=True` reduz logo/fontes/espaçamento para relatórios de uma
    única página (ex.: OPR)."""
    if LOGO_PATH.exists():
        doc.add_picture(str(LOGO_PATH), width=Cm(3.2 if compacto else 4.2))

    _paragrafo(doc, titulo, cor=_NAVY, tamanho=14 if compacto else 18, negrito=True)
    _paragrafo(doc, subtitulo, cor=_TEXTO_FRACO, tamanho=9 if compacto else 10.5)

    tamanho_meta = Pt(8.5) if compacto else Pt(10)

    p_periodo = doc.add_paragraph()
    p_periodo.add_run("Período selecionado: ").bold = True
    p_periodo.add_run(periodo_label or "Todos os períodos")
    for run in p_periodo.runs:
        run.font.size = tamanho_meta

    texto_filtros = "; ".join(f"{k}: {v}" for k, v in (filtros_aplicados or {}).items() if v)
    p_filtros = doc.add_paragraph()
    p_filtros.add_run("Filtros aplicados: ").bold = True
    p_filtros.add_run(texto_filtros or "Nenhum filtro adicional aplicado.")
    for run in p_filtros.runs:
        run.font.size = tamanho_meta

    p_geracao = doc.add_paragraph()
    p_geracao.add_run("Gerado em: ").bold = True
    p_geracao.add_run(agora_br().strftime("%d/%m/%Y às %H:%M"))
    p_geracao.add_run("    ·    Responsável: ").bold = True
    p_geracao.add_run(usuario_responsavel or "—")
    for run in p_geracao.runs:
        run.font.size = tamanho_meta

    if not compacto:
        doc.add_paragraph()


def sombrear_celula(celula, cor_hex: str) -> None:
    """Aplica uma cor de fundo (shading) a uma célula de tabela — usado
    para os indicadores de SLA nas cores padronizadas do sistema."""
    sombra = OxmlElement("w:shd")
    sombra.set(qn("w:fill"), cor_hex.lstrip("#"))
    celula._tc.get_or_add_tcPr().append(sombra)


def secao(doc: Document, titulo: str, nivel: int = 1) -> None:
    doc.add_heading(titulo, level=nivel)


def paragrafo(doc: Document, texto: str, negrito: bool = False, italico: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.bold = negrito
    run.italic = italico


def tabela_indicadores(doc: Document, pares: list[tuple[str, object]], titulo: str | None = None) -> None:
    """Tabela de duas colunas (indicador/valor) — cada indicador é uma linha
    editável da tabela nativa do Word."""
    if titulo:
        secao(doc, titulo, nivel=2)
    if not pares:
        paragrafo(doc, "Nenhum indicador disponível para os filtros aplicados.", italico=True)
        return
    tabela = doc.add_table(rows=0, cols=2)
    tabela.style = _ESTILO_TABELA
    for rotulo, valor in pares:
        linha = tabela.add_row()
        linha.cells[0].text = str(rotulo)
        linha.cells[0].paragraphs[0].runs[0].font.bold = True
        linha.cells[1].text = "—" if valor is None else str(valor)
    doc.add_paragraph()


def tabela_indicadores_compacta(doc: Document, pares: list[tuple[str, object]], titulo: str | None = None, colunas: int = 2) -> None:
    """Variante compacta da tabela de indicadores: distribui os pares
    rótulo/valor lado a lado em `colunas` blocos por linha, para caber uma
    lista mais longa de indicadores em menos espaço vertical (necessário
    para manter o OPR em uma única página)."""
    if titulo:
        secao(doc, titulo, nivel=2)
    if not pares:
        paragrafo(doc, "Nenhum indicador disponível para os filtros aplicados.", italico=True)
        return
    linhas_necessarias = -(-len(pares) // colunas)
    tabela = doc.add_table(rows=linhas_necessarias, cols=colunas * 2)
    tabela.style = _ESTILO_TABELA
    tabela.autofit = True
    for idx, (rotulo, valor) in enumerate(pares):
        linha_idx, bloco_idx = divmod(idx, colunas)
        celula_rotulo = tabela.rows[linha_idx].cells[bloco_idx * 2]
        celula_valor = tabela.rows[linha_idx].cells[bloco_idx * 2 + 1]
        celula_rotulo.text = str(rotulo)
        run_rotulo = celula_rotulo.paragraphs[0].runs[0]
        run_rotulo.font.bold = True
        run_rotulo.font.size = Pt(8.5)
        celula_valor.text = "—" if valor is None else str(valor)
        celula_valor.paragraphs[0].runs[0].font.size = Pt(8.5)
    doc.add_paragraph()


def graficos_em_grade(doc: Document, itens: list[dict], colunas: int = 2, largura_cm: float = 7.8) -> None:
    """Insere vários gráficos lado a lado em uma grade compacta (tabela do
    Word), em vez de um embaixo do outro — necessário para que o OPR
    executivo, com vários gráficos, ainda caiba em uma única página. Cada
    item é `{"fig": go.Figure, "titulo": str, "legenda": str | None}`; cada
    imagem continua individual (pode ser removida/movida/substituída)."""
    if not itens:
        return
    linhas_necessarias = -(-len(itens) // colunas)
    tabela = doc.add_table(rows=linhas_necessarias, cols=colunas)
    tabela.autofit = True
    for idx, item in enumerate(itens):
        linha_idx, coluna_idx = divmod(idx, colunas)
        celula = tabela.rows[linha_idx].cells[coluna_idx]
        p_titulo = celula.paragraphs[0]
        run_titulo = p_titulo.add_run(item["titulo"])
        run_titulo.font.bold = True
        run_titulo.font.size = Pt(8.5)
        run_titulo.font.color.rgb = _NAVY
        p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

        imagem_bytes = figura_para_imagem(item["fig"])
        p_imagem = celula.add_paragraph()
        p_imagem.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_imagem.add_run().add_picture(io.BytesIO(imagem_bytes), width=Cm(largura_cm))

        if item.get("legenda"):
            p_legenda = celula.add_paragraph()
            p_legenda.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_legenda = p_legenda.add_run(item["legenda"])
            run_legenda.italic = True
            run_legenda.font.size = Pt(7)
            run_legenda.font.color.rgb = _TEXTO_FRACO
    doc.add_paragraph()


def tabela_dataframe(doc: Document, df: pd.DataFrame, titulo: str | None = None, max_linhas: int | None = None) -> None:
    """Tabela nativa do Word a partir de um DataFrame — cabeçalho e cada
    célula continuam editáveis individualmente."""
    if titulo:
        secao(doc, titulo, nivel=2)
    if df is None or df.empty:
        paragrafo(doc, "Nenhum registro encontrado para os filtros aplicados.", italico=True)
        return
    dados = df.head(max_linhas) if max_linhas else df
    tabela = doc.add_table(rows=1, cols=len(dados.columns))
    tabela.style = _ESTILO_TABELA
    for idx, coluna in enumerate(dados.columns):
        celula = tabela.rows[0].cells[idx]
        celula.text = str(coluna)
        celula.paragraphs[0].runs[0].font.bold = True
    for _, linha in dados.iterrows():
        celulas = tabela.add_row().cells
        for idx, valor in enumerate(linha):
            celulas[idx].text = "—" if pd.isna(valor) else str(valor)
    if max_linhas and len(df) > max_linhas:
        paragrafo(doc, f"Exibindo {max_linhas} de {len(df)} registros — a base completa está disponível na tela de origem.", italico=True)
    doc.add_paragraph()


def grafico(doc: Document, fig: go.Figure, titulo: str, legenda: str | None = None, largura_cm: float = 15.5) -> None:
    """Insere um gráfico como imagem individual em alta resolução, com
    título acima e legenda abaixo — cada gráfico pode ser removido, movido
    ou substituído no Word sem afetar o restante do relatório.

    A renderização (Kaleido/Chrome) pode falhar em ambientes sem um
    Chrome/Chromium disponível — nesse caso o erro é registrado em log e
    um parágrafo informativo substitui a imagem, sem interromper a
    geração do restante do documento."""
    _paragrafo(doc, titulo, cor=_NAVY, tamanho=11.5, negrito=True)
    try:
        imagem_bytes = figura_para_imagem(fig)
    except Exception:
        _logger.exception("Falha ao renderizar o gráfico '%s' para o Word (Kaleido/Chrome indisponível).", titulo)
        paragrafo(doc, "Gráfico não disponível nesta geração do documento — os dados completos permanecem nas tabelas deste relatório.", italico=True)
        doc.add_paragraph()
        return
    doc.add_picture(io.BytesIO(imagem_bytes), width=Cm(largura_cm))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if legenda:
        p_legenda = doc.add_paragraph()
        p_legenda.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_legenda.add_run(legenda)
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = _TEXTO_FRACO
    doc.add_paragraph()


def observacoes(doc: Document, titulo: str, texto: str | None, marcador_padrao: str = "Nenhuma observação registrada para este período.") -> None:
    secao(doc, titulo, nivel=2)
    paragrafo(doc, texto or marcador_padrao)


def rodape_institucional(doc: Document) -> None:
    rodape = doc.sections[0].footer
    p = rodape.paragraphs[0] if rodape.paragraphs else rodape.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("GAT 2026 · Controle de Análises Técnicas · Tecnoplano — documento gerado automaticamente pelo sistema; totalmente editável.")
    run.font.size = Pt(8)
    run.font.color.rgb = _TEXTO_FRACO


def documento_para_bytes(doc: Document) -> bytes:
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def nome_arquivo(*partes: str) -> str:
    """Monta um nome de arquivo seguro a partir das partes informadas,
    seguindo a convenção institucional (ex.: OPR_Prestador_P0123_Julho_2026.docx)."""
    limpo = [re.sub(r"[^\w-]", "", str(p).replace(" ", "_")) for p in partes if p]
    return "_".join(limpo) + ".docx"
